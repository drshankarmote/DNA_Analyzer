from flask import Flask, render_template, request, redirect, url_for, send_file
import os
import io
import base64
import json
from matplotlib import pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches, RGBColor
from Bio import pairwise2
from Bio.pairwise2 import format_alignment
from io import BytesIO

# Force Matplotlib to use non-GUI backend
matplotlib.use('Agg')

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


# ---------- Utility Functions ----------

def count_bases(sequence):
    return {
        'A': sequence.upper().count('A'),
        'T': sequence.upper().count('T'),
        'G': sequence.upper().count('G'),
        'C': sequence.upper().count('C')
    }

def reverse_complement(sequence):
    complement = {'A': 'T', 'T': 'A', 'G': 'C', 'C': 'G'}
    return ''.join(complement.get(base, base) for base in reversed(sequence.upper()))

# Function to detect mutations
def find_mutations(original_sequence, mutated_sequence):
    alignment = pairwise2.align.globalxx(original_sequence, mutated_sequence, one_alignment_only=True)[0]
    aligned_ref = alignment.seqA
    aligned_mut = alignment.seqB

    mutations = []
    pos_ref = 0  # position in original (ungapped) reference
    for i, (ref_base, mut_base) in enumerate(zip(aligned_ref, aligned_mut)):
        if ref_base != '-':
            pos_ref += 1
        if ref_base != mut_base:
            if ref_base == '-':
                mutations.append({'type': 'insertion', 'position': pos_ref, 'inserted_base': mut_base})
            elif mut_base == '-':
                mutations.append({'type': 'deletion', 'position': pos_ref, 'deleted_base': ref_base})
            else:
                mutations.append({'type': 'substitution', 'position': pos_ref, 'original_base': ref_base, 'mutated_base': mut_base})
    return mutations

print(find_mutations)

# Codon to Amino Acid Mapping
CODON_TABLE = {
    'ATA': 'Isoleucine', 'ATC': 'Isoleucine', 'ATT': 'Isoleucine', 'ATG': 'Methionine',
    'ACA': 'Threonine', 'ACC': 'Threonine', 'ACG': 'Threonine', 'ACT': 'Threonine',
    'AAC': 'Asparagine', 'AAT': 'Asparagine', 'AAA': 'Lysine', 'AAG': 'Lysine',
    'AGC': 'Serine', 'AGT': 'Serine', 'AGA': 'Arginine', 'AGG': 'Arginine',
    'CTA': 'Leucine', 'CTC': 'Leucine', 'CTG': 'Leucine', 'CTT': 'Leucine',
    'CCA': 'Proline', 'CCC': 'Proline', 'CCG': 'Proline', 'CCT': 'Proline',
    'CAC': 'Histidine', 'CAT': 'Histidine', 'CAA': 'Glutamine', 'CAG': 'Glutamine',
    'CGA': 'Arginine', 'CGC': 'Arginine', 'CGG': 'Arginine', 'CGT': 'Arginine',
    'GTA': 'Valine', 'GTC': 'Valine', 'GTG': 'Valine', 'GTT': 'Valine',
    'GCA': 'Alanine', 'GCC': 'Alanine', 'GCG': 'Alanine', 'GCT': 'Alanine',
    'GAC': 'Aspartic acid', 'GAT': 'Aspartic acid', 'GAA': 'Glutamic acid', 'GAG': 'Glutamic acid',
    'GGA': 'Glycine', 'GGC': 'Glycine', 'GGG': 'Glycine', 'GGT': 'Glycine',
    'TCA': 'Serine', 'TCC': 'Serine', 'TCG': 'Serine', 'TCT': 'Serine',
    'TTC': 'Phenylalanine', 'TTT': 'Phenylalanine', 'TTA': 'Leucine', 'TTG': 'Leucine',
    'TAC': 'Tyrosine', 'TAT': 'Tyrosine', 'TAA': 'Stop', 'TAG': 'Stop',
    'TGC': 'Cysteine', 'TGT': 'Cysteine', 'TGA': 'Stop', 'TGG': 'Tryptophan'
}

# Function to get codons and their corresponding amino acids
def get_amino_acids(sequence):
    amino_acids = []
    for i in range(0, len(sequence), 3):  # Iterate in steps of 3 (codons)
        codon = sequence[i:i+3]
        if len(codon) == 3:
            amino_acid = CODON_TABLE.get(codon, 'X')  # 'X' for unknown codons
            amino_acids.append((codon, amino_acid))
    return amino_acids


def create_base_count_graph(base_count):
    fig, ax = plt.subplots()
    bases = list(base_count.keys())
    counts = list(base_count.values())
    colors = ['red', 'blue', 'green', 'yellow']
    ax.bar(bases, counts, color=colors)
    ax.set_ylabel('Count')
    ax.set_title('Base Count Graph')

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    graph_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
    plt.close(fig)
    return graph_b64, buf

def format_alignment_block(seq1, midline, seq2, block_size=80):
    lines = []
    for i in range(0, len(seq1), block_size):
        block_seq1 = seq1[i:i+block_size]
        block_midline = midline[i:i+block_size]
        block_seq2 = seq2[i:i+block_size]
        lines.append(block_seq1)
        lines.append(block_midline)
        lines.append(block_seq2)
        lines.append('')  # empty line between blocks
    return '\n'.join(lines)


# ---------- Routes ----------

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)

            with open(filepath, 'r') as f:
                sequence = f.read().replace("\n", "").strip()

            mutated_sequence = request.form.get('mutated_sequence', sequence)

            bases = count_bases(sequence)
            rev_comp = reverse_complement(sequence)
            graph_b64, _ = create_base_count_graph(bases)

            mutations = find_mutations(sequence, mutated_sequence)

            codons_with_amino_acids = get_amino_acids(sequence)

            return render_template('results.html', original_sequence=sequence,
                                   mutated_sequence=mutated_sequence, mutations=mutations,
                                   reverse_complement=rev_comp, base_count=bases,
                                   graph_b64=graph_b64, codons_with_amino_acids=codons_with_amino_acids)
    
    return render_template('upload.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/services')
def services():
    return render_template('services.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/download_word', methods=['POST'])
def download_word():
    sequence = request.form['original_sequence']
    rev_comp = request.form['reverse_complement']
    
    try:
        bases = json.loads(request.form['base_count'])
    except json.JSONDecodeError as e:
        print(f"JSON Decode Error: {e}")
        return "Invalid JSON data received.", 400

    graph_b64, graph_buf = create_base_count_graph(bases)

    doc = Document()
    doc.add_heading('DNA Analysis Report', level=1)

    # Original Sequence with Colors
    doc.add_heading('Original Sequence:', level=2)
    p = doc.add_paragraph()
    for base in sequence:
        run = p.add_run(base)
        if base == 'A':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif base == 'T':
            run.font.color.rgb = RGBColor(52, 152, 219)
        elif base == 'G':
            run.font.color.rgb = RGBColor(46, 204, 113)
        elif base == 'C':
            run.font.color.rgb = RGBColor(241, 196, 15)

    # Reverse Complement with Colors
    doc.add_heading('Reverse Complement:', level=2)
    p = doc.add_paragraph()
    for base in rev_comp:
        run = p.add_run(base)
        if base == 'A':
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif base == 'T':
            run.font.color.rgb = RGBColor(52, 152, 219)
        elif base == 'G':
            run.font.color.rgb = RGBColor(46, 204, 113)
        elif base == 'C':
            run.font.color.rgb = RGBColor(241, 196, 15)

    # Base Count
    doc.add_heading('Base Counts:', level=2)
    for base, count in bases.items():
        doc.add_paragraph(f"{base}: {count}")

    # Graph
    img_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_graph.png')
    with open(img_path, 'wb') as f:
        f.write(graph_buf.getbuffer())
    doc.add_picture(img_path, width=Inches(5))
    os.remove(img_path)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="DNA_Report.docx")

@app.route('/compare_sequences', methods=['POST'])
def compare_sequences():
    seq1 = request.form['seq1'].upper().replace("\n", "").strip()
    seq2 = request.form['seq2'].upper().replace("\n", "").strip()

    comparison = []
    match_count = 0
    mismatch_count = 0

    for i in range(min(len(seq1), len(seq2))):
        base1 = seq1[i]
        base2 = seq2[i]
        if base1 == base2:
            comparison.append({'index': i + 1, 'base1': base1, 'base2': base2, 'result': '✔ Match'})
            match_count += 1
        else:
            comparison.append({'index': i + 1, 'base1': base1, 'base2': base2, 'result': '✘ Mismatch'})
            mismatch_count += 1

    similarity_percentage = round((match_count / min(len(seq1), len(seq2))) * 100, 2)

    # Global alignment using Biopython
    alignment = pairwise2.align.globalxx(seq1, seq2, one_alignment_only=True)[0]
    aligned_seq1 = alignment.seqA
    aligned_seq2 = alignment.seqB
    alignment_score = alignment.score
    midline = ''.join('|' if a == b else ' ' for a, b in zip(aligned_seq1, aligned_seq2))
    formatted_alignment = format_alignment_block(aligned_seq1, midline, aligned_seq2)

    return render_template('comparison_result.html',
                           seq1=seq1,
                           seq2=seq2,
                           comparison=comparison,
                           match_count=match_count,
                           mismatch_count=mismatch_count,
                           similarity_percentage=similarity_percentage,
                           alignment_result=formatted_alignment,
                           alignment_score=alignment_score)


# ---------- Server Start ----------

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)
